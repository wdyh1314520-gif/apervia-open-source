# Split from app3_parts/tools/file_registry_edit_tools_part.py.
# Purpose: sandbox Office generation helpers.
# Loaded by app3.py via _exec_split_file(...), sharing the original global namespace.

def _sandbox_office_limit_args(args: dict | None = None) -> tuple[bool, str, int]:
    raw = json.dumps(args or {}, ensure_ascii=False, default=str)
    try:
        max_chars = max(10000, min(int(app_getenv('SANDBOX_OFFICE_MAX_JSON_CHARS', str(2 * 1024 * 1024)) or (2 * 1024 * 1024)), 20 * 1024 * 1024))
    except Exception:
        max_chars = 2 * 1024 * 1024
    if len(raw) > max_chars:
        return False, 'office_payload_too_large', max_chars
    return True, '', max_chars


def _sandbox_office_text(value, limit: int = 20000) -> str:
    text = str(value if value is not None else '')
    text = text.replace('\x00', '')
    if len(text) > max(1, int(limit or 1)):
        text = text[:max(1, int(limit or 1))]
    return text


def _sandbox_office_rows(headers=None, rows=None, max_rows: int = 2000, max_cols: int = 50) -> list[list[str]]:
    out: list[list[str]] = []
    header_list = headers if isinstance(headers, list) else []
    if header_list:
        out.append([_sandbox_office_text(x, 4000) for x in header_list[:max_cols]])
    source_rows = rows if isinstance(rows, list) else []
    for row in source_rows[:max_rows]:
        if isinstance(row, dict):
            keys = header_list or list(row.keys())
            out.append([_sandbox_office_text(row.get(k), 4000) for k in keys[:max_cols]])
        elif isinstance(row, (list, tuple)):
            out.append([_sandbox_office_text(x, 4000) for x in list(row)[:max_cols]])
        else:
            out.append([_sandbox_office_text(row, 4000)])
    return out


def _sandbox_office_markdown_table_cells(line: str = '') -> list[str]:
    """Return Markdown pipe-table cells, or [] when the line is not table-like.

    This intentionally stays lightweight: it handles the normal GFM form
    used by LLMs (| A | B |) without pulling in a full Markdown parser, so
    Office generation remains dependency-free and predictable.
    """
    raw = str(line or '').strip()
    if '|' not in raw:
        return []
    # Avoid treating shell pipes as tables. Real Markdown tables normally have
    # multiple pipe separators or a leading/trailing pipe.
    if raw.count('|') < 2 and not (raw.startswith('|') or raw.endswith('|')):
        return []
    if raw.startswith('|'):
        raw = raw[1:]
    if raw.endswith('|'):
        raw = raw[:-1]
    cells = [c.strip() for c in raw.split('|')]
    if len(cells) < 2:
        return []
    return cells


def _sandbox_office_markdown_is_table_separator(line: str = '') -> bool:
    cells = _sandbox_office_markdown_table_cells(line)
    if len(cells) < 2:
        return False
    ok = 0
    for cell in cells:
        token = re.sub(r'\s+', '', str(cell or ''))
        if re.match(r'^:?-{3,}:?$', token):
            ok += 1
    return ok == len(cells)


def _sandbox_office_markdown_is_table_start(lines: list[str], index: int) -> bool:
    if index < 0 or index + 1 >= len(lines):
        return False
    header = _sandbox_office_markdown_table_cells(lines[index])
    if len(header) < 2:
        return False
    if not _sandbox_office_markdown_is_table_separator(lines[index + 1]):
        return False
    sep_cells = _sandbox_office_markdown_table_cells(lines[index + 1])
    return len(sep_cells) >= 2


def _sandbox_office_markdown_normalize_table_row(cells: list[str], width: int) -> list[str]:
    width = max(1, int(width or 1))
    row = [_sandbox_office_text(c, 4000) for c in list(cells or [])[:width]]
    while len(row) < width:
        row.append('')
    return row


def _sandbox_office_markdown_looks_structured(text: str = '') -> bool:
    raw = str(text or '').replace('\r\n', '\n').replace('\r', '\n')
    if not raw.strip():
        return False
    signals = 0
    if re.search(r'(?m)^\s*#{1,6}\s+\S+', raw):
        signals += 1
    if re.search(r'(?m)^\s*```', raw):
        signals += 2
    if re.search(r'(?m)^\s*[-*]\s+\S+', raw) or re.search(r'(?m)^\s*\d+[\.)]\s+\S+', raw):
        signals += 1
    if re.search(r'(?m)^\s*>\s+\S+', raw):
        signals += 1
    lines = raw.split('\n')
    if any(_sandbox_office_markdown_is_table_start(lines, idx) for idx in range(0, max(0, len(lines) - 1))):
        signals += 2
    return signals >= 2


def _sandbox_office_markdown_line_signal(text: str = '') -> bool:
    stripped = str(text or '').strip()
    if not stripped:
        return False
    if re.match(r'^```\s*[A-Za-z0-9_+.-]*\s*$', stripped):
        return True
    if _sandbox_office_markdown_is_table_separator(stripped):
        return True
    if len(_sandbox_office_markdown_table_cells(stripped)) >= 2 and stripped.startswith('|'):
        return True
    if re.match(r'^#{1,6}\s+\S+', stripped):
        return True
    if re.match(r'^>\s+\S+', stripped):
        return True
    if re.match(r'^[-*]\s+\S+', stripped) or re.match(r'^\d+[\.)]\s+\S+', stripped):
        return True
    return False


def _sandbox_office_markdown_sections(text: str = '') -> list[dict]:
    raw = _sandbox_office_text(text, 500000).replace('\r\n', '\n').replace('\r', '\n')
    lines = raw.split('\n')
    out: list[dict] = []
    para: list[str] = []
    list_items: list[str] = []
    list_ordered = False
    quote: list[str] = []
    code: list[str] = []
    in_code = False

    def flush_para() -> None:
        nonlocal para
        joined = '\n'.join(x.strip() for x in para if str(x or '').strip()).strip()
        para = []
        if joined:
            out.append({'type': 'paragraph', 'text': joined})

    def flush_list() -> None:
        nonlocal list_items, list_ordered
        if list_items:
            out.append({'type': 'list', 'ordered': bool(list_ordered), 'items': list_items[:]})
        list_items = []
        list_ordered = False

    def flush_quote() -> None:
        nonlocal quote
        joined = '\n'.join(x.strip() for x in quote if str(x or '').strip()).strip()
        quote = []
        if joined:
            out.append({'type': 'note', 'text': joined})

    def flush_code() -> None:
        nonlocal code
        joined = '\n'.join(code).strip('\n')
        code = []
        if joined:
            out.append({'type': 'code', 'text': joined})

    idx = 0
    while idx < len(lines):
        raw_line = lines[idx]
        line = str(raw_line or '').rstrip()
        stripped = line.strip()
        fence = re.match(r'^```\s*([A-Za-z0-9_+.-]*)\s*$', stripped)
        if fence:
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_para()
                flush_list()
                flush_quote()
                in_code = True
                code = []
            idx += 1
            continue
        if in_code:
            code.append(line)
            idx += 1
            continue
        if _sandbox_office_markdown_is_table_start(lines, idx):
            flush_para()
            flush_list()
            flush_quote()
            headers = _sandbox_office_markdown_table_cells(lines[idx])
            width = max(2, len(headers))
            rows: list[list[str]] = []
            idx += 2
            while idx < len(lines):
                row_line = str(lines[idx] or '').rstrip()
                row_stripped = row_line.strip()
                if not row_stripped:
                    break
                if _sandbox_office_markdown_is_table_separator(row_line):
                    idx += 1
                    continue
                cells = _sandbox_office_markdown_table_cells(row_line)
                if len(cells) < 2:
                    break
                width = max(width, len(cells))
                rows.append(cells)
                idx += 1
            out.append({
                'type': 'table',
                'headers': _sandbox_office_markdown_normalize_table_row(headers, width),
                'rows': [_sandbox_office_markdown_normalize_table_row(r, width) for r in rows],
            })
            continue
        if not stripped:
            flush_para()
            flush_list()
            flush_quote()
            idx += 1
            continue
        if re.match(r'^-{3,}$', stripped):
            flush_para()
            flush_list()
            flush_quote()
            idx += 1
            continue
        heading = re.match(r'^(#{1,6})\s+(.+?)\s*$', stripped)
        if heading:
            flush_para()
            flush_list()
            flush_quote()
            out.append({'type': 'heading', 'level': min(6, len(heading.group(1))), 'text': heading.group(2).strip()})
            idx += 1
            continue
        quote_match = re.match(r'^>\s?(.*)$', stripped)
        if quote_match:
            flush_para()
            flush_list()
            quote.append(quote_match.group(1).strip())
            idx += 1
            continue
        bullet_match = re.match(r'^[-*]\s+(.+)$', stripped)
        number_match = re.match(r'^\d+[\.)]\s+(.+)$', stripped)
        if bullet_match or number_match:
            flush_para()
            flush_quote()
            ordered = bool(number_match)
            if list_items and list_ordered != ordered:
                flush_list()
            list_ordered = ordered
            list_items.append((number_match.group(1) if number_match else bullet_match.group(1)).strip())
            idx += 1
            continue
        flush_list()
        flush_quote()
        para.append(stripped)
        idx += 1
    if in_code:
        flush_code()
    flush_para()
    flush_list()
    flush_quote()
    return out[:800]


def _sandbox_office_sections(args: dict | None = None) -> list[dict]:
    args = dict(args or {})
    sections = args.get('sections')
    if not isinstance(sections, list):
        sections = []
    out = []

    pending_markdown_lines: list[str] = []

    def flush_pending_markdown_lines() -> None:
        nonlocal pending_markdown_lines
        if not pending_markdown_lines:
            return
        raw = '\n'.join(str(x or '') for x in pending_markdown_lines).strip()
        pending_markdown_lines = []
        if not raw:
            return
        if _sandbox_office_markdown_looks_structured(raw):
            out.extend(_sandbox_office_markdown_sections(raw))
        else:
            for line in raw.split('\n'):
                line = str(line or '').strip()
                if line:
                    out.append({'type': 'paragraph', 'text': line})

    for section in [x for x in sections if isinstance(x, dict)][:800]:
        row = dict(section)
        typ = str(row.get('type') or row.get('kind') or '').strip().lower()
        raw_text = _sandbox_office_text(row.get('content') if row.get('content') not in (None, '') else row.get('text'), 500000)
        if typ in {'', 'paragraph', 'text', 'body'} and raw_text:
            if pending_markdown_lines or _sandbox_office_markdown_line_signal(raw_text) or _sandbox_office_markdown_looks_structured(raw_text):
                pending_markdown_lines.extend(raw_text.replace('\r\n', '\n').replace('\r', '\n').split('\n'))
                continue
            flush_pending_markdown_lines()
            out.append(row)
        else:
            flush_pending_markdown_lines()
            out.append(row)
    flush_pending_markdown_lines()
    content = _sandbox_office_text(args.get('content'), 500000)
    if content and not out:
        if _sandbox_office_markdown_looks_structured(content):
            out.extend(_sandbox_office_markdown_sections(content))
        else:
            for para in re.split(r'\n\s*\n', content)[:160]:
                para = para.strip()
                if para:
                    out.append({'type': 'paragraph', 'text': para})
    title = _sandbox_office_text(args.get('title'), 500)
    first_text = str((out[0] or {}).get('text') or (out[0] or {}).get('title') or '').strip() if out else ''
    if title and (not out or str(out[0].get('type') or '').lower() != 'heading') and first_text.strip().lower() != title.strip().lower():
        out.insert(0, {'type': 'heading', 'level': 1, 'text': title})
    return out[:800]


def _sandbox_office_xml_escape(value) -> str:
    return (
        _sandbox_office_text(value, 20000)
        .replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
        .replace('"', '&quot;')
        .replace("'", '&apos;')
    )


def _sandbox_docx_set_style_font(style, *, font_name: str = 'Calibri', east_asia: str = 'Microsoft YaHei', size_pt: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    try:
        from docx.shared import Pt, RGBColor  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        font = style.font
        font.name = font_name
        if size_pt is not None:
            font.size = Pt(float(size_pt))
        if bold is not None:
            font.bold = bool(bold)
        if color:
            color = str(color).strip().lstrip('#')
            if len(color) == 6:
                font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = rpr._add_rFonts()
        rfonts.set(qn('w:eastAsia'), east_asia)
    except Exception:
        pass


def _sandbox_docx_shade_paragraph(paragraph, fill: str = 'F4F6F9') -> None:
    try:
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        ppr = paragraph._p.get_or_add_pPr()
        shd = ppr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd')
            ppr.append(shd)
        shd.set(qn('w:fill'), str(fill).strip().lstrip('#')[:6] or 'F4F6F9')
    except Exception:
        pass


def _sandbox_docx_set_cell_margins(cell, top: int = 100, right: int = 120, bottom: int = 100, left: int = 120) -> None:
    try:
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.find(qn('w:tcMar'))
        if tc_mar is None:
            tc_mar = OxmlElement('w:tcMar')
            tc_pr.append(tc_mar)
        for name, value in [('top', top), ('right', right), ('bottom', bottom), ('left', left)]:
            node = tc_mar.find(qn(f'w:{name}'))
            if node is None:
                node = OxmlElement(f'w:{name}')
                tc_mar.append(node)
            node.set(qn('w:w'), str(max(0, int(value or 0))))
            node.set(qn('w:type'), 'dxa')
    except Exception:
        pass


def _sandbox_docx_shade_cell(cell, fill: str = 'EAF2F8') -> None:
    try:
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd')
            tc_pr.append(shd)
        shd.set(qn('w:fill'), str(fill).strip().lstrip('#')[:6] or 'EAF2F8')
    except Exception:
        pass


def _sandbox_docx_table_widths(rows: list[list[str]], usable_width_dxa: int = 9360) -> list[int]:
    max_cols = max((len(r) for r in rows), default=1)
    weights = []
    for c_idx in range(max_cols):
        max_len = 0
        for row in rows[:120]:
            text = str(row[c_idx] if c_idx < len(row) else '')
            max_len = max(max_len, min(len(text), 80))
        weights.append(max(8, min(max_len, 42)))
    total = max(1, sum(weights))
    min_col = 900 if max_cols <= 6 else 650
    widths = [max(min_col, int(usable_width_dxa * (w / total))) for w in weights]
    overflow = sum(widths) - usable_width_dxa
    while overflow > 0 and any(w > min_col for w in widths):
        adjustable = [idx for idx, w in enumerate(widths) if w > min_col]
        step = max(1, int(overflow / max(1, len(adjustable))))
        for idx in adjustable:
            take = min(step, widths[idx] - min_col, overflow)
            widths[idx] -= take
            overflow -= take
            if overflow <= 0:
                break
    return widths


def _sandbox_docx_apply_table_geometry(table, rows: list[list[str]], usable_width_dxa: int = 9360) -> None:
    try:
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Pt  # type: ignore
        widths = _sandbox_docx_table_widths(rows, usable_width_dxa=usable_width_dxa)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        tbl = table._tbl
        tbl_pr = table._tbl.tblPr
        tbl_ind = tbl_pr.find(qn('w:tblInd'))
        if tbl_ind is None:
            tbl_ind = OxmlElement('w:tblInd')
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn('w:w'), '120')
        tbl_ind.set(qn('w:type'), 'dxa')
        tbl_w = tbl_pr.find(qn('w:tblW'))
        if tbl_w is None:
            tbl_w = OxmlElement('w:tblW')
            tbl_pr.append(tbl_w)
        tbl_w.set(qn('w:w'), str(sum(widths)))
        tbl_w.set(qn('w:type'), 'dxa')
        tbl_layout = tbl_pr.find(qn('w:tblLayout'))
        if tbl_layout is None:
            tbl_layout = OxmlElement('w:tblLayout')
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn('w:type'), 'fixed')
        tbl_grid = tbl.tblGrid
        if tbl_grid is None:
            tbl_grid = OxmlElement('w:tblGrid')
            tbl.insert(1, tbl_grid)
        for child in list(tbl_grid):
            tbl_grid.remove(child)
        for width in widths:
            grid_col = OxmlElement('w:gridCol')
            grid_col.set(qn('w:w'), str(width))
            tbl_grid.append(grid_col)
        for col_idx, column in enumerate(table.columns):
            column.width = widths[min(col_idx, len(widths) - 1)]
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                width = widths[min(col_idx, len(widths) - 1)]
                cell.width = width
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _sandbox_docx_set_cell_margins(cell)
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn('w:tcW'))
                if tc_w is None:
                    tc_w = OxmlElement('w:tcW')
                    tc_pr.append(tc_w)
                tc_w.set(qn('w:w'), str(width))
                tc_w.set(qn('w:type'), 'dxa')
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = 0
                    para.paragraph_format.space_after = 0
                    para.paragraph_format.line_spacing = 1.08
                    for run in para.runs:
                        run.font.size = Pt(9.5)
                        if row_idx == 0:
                            run.font.bold = True
                if row_idx == 0:
                    _sandbox_docx_shade_cell(cell, 'EAF2F8')
    except Exception:
        pass


def _sandbox_prepare_docx_styles(doc) -> None:
    try:
        from docx.enum.style import WD_STYLE_TYPE  # type: ignore
        from docx.shared import Pt, Inches  # type: ignore
        styles = doc.styles
        _sandbox_docx_set_style_font(styles['Normal'], font_name='Calibri', east_asia='Microsoft YaHei', size_pt=10.5)
        styles['Normal'].paragraph_format.space_after = Pt(6)
        styles['Normal'].paragraph_format.line_spacing = 1.15
        for name, size, color in [('Heading 1', 16, '2E74B5'), ('Heading 2', 13, '2E74B5'), ('Heading 3', 12, '1F4D78')]:
            if name in styles:
                _sandbox_docx_set_style_font(styles[name], font_name='Calibri', east_asia='Microsoft YaHei', size_pt=size, bold=True, color=color)
                styles[name].paragraph_format.space_before = Pt(10 if name == 'Heading 1' else 8)
                styles[name].paragraph_format.space_after = Pt(6 if name == 'Heading 1' else 4)
        for name in ('List Number', 'List Bullet'):
            if name in styles:
                _sandbox_docx_set_style_font(styles[name], font_name='Calibri', east_asia='Microsoft YaHei', size_pt=10.5)
                styles[name].paragraph_format.space_after = Pt(4)
        if 'Sandbox Code' not in styles:
            code_style = styles.add_style('Sandbox Code', WD_STYLE_TYPE.PARAGRAPH)
        else:
            code_style = styles['Sandbox Code']
        _sandbox_docx_set_style_font(code_style, font_name='Consolas', east_asia='Microsoft YaHei', size_pt=9.5, color='1F2937')
        code_style.paragraph_format.left_indent = Inches(0.18)
        code_style.paragraph_format.right_indent = Inches(0.05)
        code_style.paragraph_format.space_before = Pt(2)
        code_style.paragraph_format.space_after = Pt(2)
        if 'Sandbox Note' not in styles:
            note_style = styles.add_style('Sandbox Note', WD_STYLE_TYPE.PARAGRAPH)
        else:
            note_style = styles['Sandbox Note']
        _sandbox_docx_set_style_font(note_style, font_name='Calibri', east_asia='Microsoft YaHei', size_pt=10.5, color='374151')
        note_style.paragraph_format.left_indent = Inches(0.12)
        note_style.paragraph_format.space_before = Pt(4)
        note_style.paragraph_format.space_after = Pt(6)
    except Exception:
        pass


def _sandbox_docx_is_command_line(text: str = '') -> bool:
    text = str(text or '').strip()
    if not text or len(text) > 240:
        return False
    if re.match(r'^(sudo|curl|wget|bash|sh|systemctl|service|docker|apt|apt-get|dnf|yum|ufw|iptables|firewall-cmd|certbot|openssl|ssh|scp|rsync|chmod|chown|mkdir|cp|mv|cat|nano|vim|vi|grep|sed|awk)\b', text, flags=re.I):
        return True
    if re.match(r'^(~?/|\./|[A-Za-z]:[\\/])', text):
        return True
    if any(mark in text for mark in (' && ', ' || ', ' | ', ' > ', ' <(')):
        return True
    return False


def _sandbox_docx_add_inline_runs(paragraph, text: str = '') -> None:
    raw = str(text or '')
    if not raw:
        return
    try:
        from docx.shared import RGBColor  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
    except Exception:
        RGBColor = None
        qn = None
    pattern = re.compile(r'(`[^`\n]+`|\*\*[^*\n]+\*\*)')
    pos = 0

    def add_run(value: str = '', *, code: bool = False, bold: bool = False) -> None:
        if not value:
            return
        run = paragraph.add_run(value)
        if bold:
            run.bold = True
        if code:
            try:
                run.font.name = 'Consolas'
                if qn is not None:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                if RGBColor is not None:
                    run.font.color.rgb = RGBColor(31, 41, 55)
            except Exception:
                pass

    for match in pattern.finditer(raw):
        add_run(raw[pos:match.start()])
        token = match.group(0)
        if token.startswith('`') and token.endswith('`'):
            add_run(token[1:-1], code=True)
        elif token.startswith('**') and token.endswith('**'):
            add_run(token[2:-2], bold=True)
        else:
            add_run(token)
        pos = match.end()
    add_run(raw[pos:])


def _sandbox_docx_add_paragraph_block(doc, text: str = '') -> None:
    text = _sandbox_office_text(text, 50000).replace('\r\n', '\n').replace('\r', '\n')
    if not text.strip():
        return
    if _sandbox_office_markdown_looks_structured(text):
        for section in _sandbox_office_markdown_sections(text):
            typ = str(section.get('type') or section.get('kind') or 'paragraph').strip().lower()
            if typ in {'heading', 'title'}:
                try:
                    level = max(1, min(int(section.get('level') or 1), 6))
                except Exception:
                    level = 1
                heading = doc.add_heading('', level=level)
                _sandbox_docx_add_inline_runs(heading, _sandbox_office_text(section.get('text') or section.get('title'), 4000))
                try:
                    heading.paragraph_format.keep_with_next = True
                except Exception:
                    pass
            elif typ in {'list', 'bullets', 'bullet_list'}:
                style = 'List Number' if bool(section.get('ordered') or section.get('numbered')) else 'List Bullet'
                for item in (section.get('items') if isinstance(section.get('items'), list) else []):
                    para = doc.add_paragraph(style=style)
                    _sandbox_docx_add_inline_runs(para, _sandbox_office_text(item, 4000))
            elif typ in {'code', 'code_block', 'command', 'commands'}:
                for line in _sandbox_office_text(section.get('text') or section.get('content'), 50000).split('\n'):
                    if str(line).strip():
                        para = doc.add_paragraph(str(line).strip(), style='Sandbox Code')
                        _sandbox_docx_shade_paragraph(para, 'F3F4F6')
            elif typ in {'note', 'callout', 'warning'}:
                para = doc.add_paragraph(style='Sandbox Note')
                _sandbox_docx_add_inline_runs(para, _sandbox_office_text(section.get('text') or section.get('content'), 20000))
                _sandbox_docx_shade_paragraph(para, 'F4F6F9' if typ != 'warning' else 'FFF7ED')
            elif typ == 'table':
                rows = _sandbox_office_rows(section.get('headers'), section.get('rows'), max_rows=400, max_cols=24)
                if rows:
                    table = doc.add_table(rows=len(rows), cols=max(1, max(len(r) for r in rows)))
                    table.style = 'Table Grid'
                    for r_idx, row in enumerate(rows):
                        for c_idx, cell_text in enumerate(row):
                            table.rows[r_idx].cells[c_idx].text = cell_text
                    _sandbox_docx_apply_table_geometry(table, rows, usable_width_dxa=9360)
            else:
                _sandbox_docx_add_paragraph_block(doc, _sandbox_office_text(section.get('text') or section.get('content'), 50000))
        return
    lines = text.split('\n')
    normal_buffer: list[str] = []

    def flush_normal() -> None:
        nonlocal normal_buffer
        joined = '\n'.join(x for x in normal_buffer if str(x or '').strip()).strip()
        normal_buffer = []
        if joined:
            para = doc.add_paragraph()
            _sandbox_docx_add_inline_runs(para, joined)

    for raw_line in lines:
        line = str(raw_line or '').strip()
        if not line:
            flush_normal()
            continue
        number_match = re.match(r'^\s*\d+[\.)、]\s+(.+)$', line)
        bullet_match = re.match(r'^\s*[-*•]\s+(.+)$', line)
        if number_match:
            flush_normal()
            para = doc.add_paragraph(style='List Number')
            _sandbox_docx_add_inline_runs(para, number_match.group(1).strip())
        elif bullet_match:
            flush_normal()
            para = doc.add_paragraph(style='List Bullet')
            _sandbox_docx_add_inline_runs(para, bullet_match.group(1).strip())
        elif _sandbox_docx_is_command_line(line):
            flush_normal()
            para = doc.add_paragraph(line, style='Sandbox Code')
            _sandbox_docx_shade_paragraph(para, 'F3F4F6')
        else:
            normal_buffer.append(line)
    flush_normal()


def _sandbox_docx_structure_diagnostics(raw: bytes) -> dict:
    info = {
        'validator': 'python-docx',
        'valid': True,
        'real_table_count': 0,
        'markdown_table_like_paragraphs': 0,
        'markdown_fence_like_paragraphs': 0,
    }
    try:
        import zipfile
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        with zipfile.ZipFile(io.BytesIO(raw), 'r') as zf:
            root = ET.fromstring(zf.read('word/document.xml'))
        info['real_table_count'] = len(root.findall('.//w:tbl', ns))
        for p in root.findall('.//w:p', ns):
            text = ''.join((t.text or '') for t in p.findall('.//w:t', ns)).strip()
            if not text:
                continue
            if text.startswith('|') and len(_sandbox_office_markdown_table_cells(text)) >= 2:
                info['markdown_table_like_paragraphs'] += 1
            if text.startswith('```') or '```' in text:
                info['markdown_fence_like_paragraphs'] += 1
        if info['markdown_table_like_paragraphs'] or info['markdown_fence_like_paragraphs']:
            info['valid'] = False
            info['warning'] = 'docx_contains_markdown_table_or_fence_text'
    except Exception as exc:
        info['valid'] = False
        info['warning'] = f'{type(exc).__name__}: {exc}'
    return info


def _sandbox_make_docx_bytes(args: dict | None = None) -> tuple[bytes, dict]:
    args = dict(args or {})
    Doc = globals().get('Document')
    if not callable(Doc):
        from docx import Document as Doc  # type: ignore
    doc = Doc()
    usable_width_dxa = 9360
    try:
        from docx.shared import Inches  # type: ignore
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    except Exception:
        pass
    _sandbox_prepare_docx_styles(doc)
    for section in _sandbox_office_sections(args):
        typ = str(section.get('type') or section.get('kind') or 'paragraph').strip().lower()
        if typ in {'heading', 'title'}:
            try:
                level = max(1, min(int(section.get('level') or 1), 6))
            except Exception:
                level = 1
            heading = doc.add_heading('', level=level)
            _sandbox_docx_add_inline_runs(heading, _sandbox_office_text(section.get('text') or section.get('title'), 4000))
            try:
                heading.paragraph_format.keep_with_next = True
            except Exception:
                pass
        elif typ in {'list', 'bullets', 'bullet_list'}:
            items = section.get('items') if isinstance(section.get('items'), list) else section.get('bullets')
            style = 'List Number' if bool(section.get('ordered') or section.get('numbered')) else 'List Bullet'
            for item in (items if isinstance(items, list) else []):
                para = doc.add_paragraph(style=style)
                _sandbox_docx_add_inline_runs(para, _sandbox_office_text(item, 4000))
        elif typ in {'code', 'code_block', 'command', 'commands'}:
            raw_lines = section.get('lines') if isinstance(section.get('lines'), list) else None
            code_text = '\n'.join(_sandbox_office_text(x, 4000) for x in raw_lines) if raw_lines is not None else _sandbox_office_text(section.get('text') or section.get('content'), 50000)
            for line in code_text.replace('\r\n', '\n').replace('\r', '\n').split('\n'):
                if str(line).strip():
                    para = doc.add_paragraph(str(line).strip(), style='Sandbox Code')
                    _sandbox_docx_shade_paragraph(para, 'F3F4F6')
        elif typ in {'note', 'callout', 'warning'}:
            para = doc.add_paragraph(style='Sandbox Note')
            _sandbox_docx_add_inline_runs(para, _sandbox_office_text(section.get('text') or section.get('content'), 20000))
            _sandbox_docx_shade_paragraph(para, 'F4F6F9' if typ != 'warning' else 'FFF7ED')
        elif typ == 'table':
            rows = _sandbox_office_rows(section.get('headers'), section.get('rows'), max_rows=400, max_cols=24)
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=max(1, max(len(r) for r in rows)))
            table.style = 'Table Grid'
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    table.rows[r_idx].cells[c_idx].text = cell_text
            _sandbox_docx_apply_table_geometry(table, rows, usable_width_dxa=usable_width_dxa)
        elif typ in {'page_break', 'pagebreak'}:
            doc.add_page_break()
        else:
            _sandbox_docx_add_paragraph_block(doc, _sandbox_office_text(section.get('text') or section.get('content'), 50000))
    bio = io.BytesIO()
    doc.save(bio)
    raw = bio.getvalue()
    Doc(io.BytesIO(raw))
    return raw, _sandbox_docx_structure_diagnostics(raw)


def _sandbox_make_xlsx_bytes(args: dict | None = None) -> tuple[bytes, dict]:
    args = dict(args or {})
    wb = openpyxl.Workbook()
    default = wb.active
    sheets = args.get('sheets')
    if not isinstance(sheets, list) or not sheets:
        tables = [s for s in _sandbox_office_sections(args) if str(s.get('type') or '').lower() == 'table']
        if tables:
            sheets = [{'name': 'Sheet1', 'headers': tables[0].get('headers'), 'rows': tables[0].get('rows')}]
        else:
            sheets = [{'name': 'Sheet1', 'headers': ['Content'], 'rows': [[_sandbox_office_text(args.get('content') or args.get('title') or '', 100000)]]}]
    used_default = False
    total_rows = 0
    for idx, sheet in enumerate([x for x in sheets if isinstance(x, dict)][:30]):
        name = _sandbox_office_text(sheet.get('name') or f'Sheet{idx + 1}', 31).strip() or f'Sheet{idx + 1}'
        name = re.sub(r'[\[\]\:\*\?\/\\]', '_', name)[:31] or f'Sheet{idx + 1}'
        ws = default if not used_default else wb.create_sheet(title=name)
        used_default = True
        ws.title = name
        rows = _sandbox_office_rows(sheet.get('headers'), sheet.get('rows'), max_rows=3000, max_cols=80)
        for row in rows:
            ws.append(row)
            total_rows += 1
        if rows and sheet.get('headers'):
            ws.freeze_panes = 'A2'
    bio = io.BytesIO()
    wb.save(bio)
    raw = bio.getvalue()
    openpyxl.load_workbook(io.BytesIO(raw), data_only=False)
    return raw, {'validator': 'openpyxl', 'valid': True, 'rows': total_rows}


def _sandbox_make_pdf_bytes(args: dict | None = None) -> tuple[bytes, dict]:
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.pdfbase import pdfmetrics  # type: ignore
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont  # type: ignore
    from reportlab.pdfgen import canvas  # type: ignore
    args = dict(args or {})
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4
    margin = 54
    y = height - margin
    font_name = 'Helvetica'
    try:
        pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
        font_name = 'STSong-Light'
    except Exception:
        font_name = 'Helvetica'

    def draw_line(text: str = '', size: int = 11, leading: int = 15):
        nonlocal y
        if y < margin + leading:
            c.showPage()
            y = height - margin
        c.setFont(('Helvetica-Bold' if size >= 15 else 'Helvetica') if font_name == 'Helvetica' else font_name, size)
        max_chars = max(24, int((width - margin * 2) / max(size * 0.52, 5)))
        for chunk in re.findall(r'.{1,' + str(max_chars) + r'}(?:\s+|$)', text) or [text]:
            if y < margin + leading:
                c.showPage()
                y = height - margin
                c.setFont(('Helvetica-Bold' if size >= 15 else 'Helvetica') if font_name == 'Helvetica' else font_name, size)
            c.drawString(margin, y, chunk.strip())
            y -= leading

    for section in _sandbox_office_sections(args):
        typ = str(section.get('type') or 'paragraph').strip().lower()
        if typ in {'heading', 'title'}:
            draw_line(_sandbox_office_text(section.get('text') or section.get('title'), 4000), 16, 22)
        elif typ in {'list', 'bullets', 'bullet_list'}:
            items = section.get('items') if isinstance(section.get('items'), list) else section.get('bullets')
            for item in (items if isinstance(items, list) else []):
                draw_line('- ' + _sandbox_office_text(item, 4000), 11, 15)
        elif typ == 'table':
            for row in _sandbox_office_rows(section.get('headers'), section.get('rows'), max_rows=120, max_cols=8):
                draw_line(' | '.join(row), 9, 13)
        elif typ in {'page_break', 'pagebreak'}:
            c.showPage()
            y = height - margin
        else:
            for para in _sandbox_office_text(section.get('text') or section.get('content'), 20000).splitlines() or ['']:
                draw_line(para, 11, 15)
            y -= 6
    c.save()
    raw = bio.getvalue()
    valid = raw.startswith(b'%PDF') and b'%%EOF' in raw[-2048:]
    if not valid:
        raise ValueError('pdf_validation_failed')
    return raw, {'validator': 'reportlab_pdf_header', 'valid': True}


def _sandbox_make_pptx_bytes(args: dict | None = None) -> tuple[bytes, dict]:
    args = dict(args or {})
    slides = args.get('slides')
    if not isinstance(slides, list) or not slides:
        slides = []
        current = {'title': _sandbox_office_text(args.get('title') or 'Presentation', 400), 'bullets': []}
        for section in _sandbox_office_sections(args):
            typ = str(section.get('type') or '').lower()
            if typ in {'heading', 'title'} and current.get('bullets'):
                slides.append(current)
                current = {'title': _sandbox_office_text(section.get('text') or section.get('title'), 400), 'bullets': []}
            elif typ in {'heading', 'title'}:
                current['title'] = _sandbox_office_text(section.get('text') or section.get('title'), 400)
            elif typ in {'list', 'bullets', 'bullet_list'}:
                items = section.get('items') if isinstance(section.get('items'), list) else section.get('bullets')
                current.setdefault('bullets', []).extend([_sandbox_office_text(x, 1000) for x in (items if isinstance(items, list) else [])])
            elif typ == 'paragraph':
                txt = _sandbox_office_text(section.get('text') or section.get('content'), 1000)
                if txt:
                    current.setdefault('bullets', []).append(txt)
        slides.append(current)
    slides = [x for x in slides if isinstance(x, dict)][:80] or [{'title': 'Presentation', 'bullets': []}]

    def rels(rows: list[tuple[str, str, str]]) -> str:
        body = ''.join(f'<Relationship Id="{_sandbox_office_xml_escape(rid)}" Type="{_sandbox_office_xml_escape(typ)}" Target="{_sandbox_office_xml_escape(target)}"/>' for rid, typ, target in rows)
        return '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">' + body + '</Relationships>'

    def text_shape(shape_id: int, name: str, x: int, y: int, cx: int, cy: int, lines: list[str], font_size: int = 2400, bold: bool = False) -> str:
        paras = []
        for line in lines or ['']:
            paras.append(
                '<a:p><a:r><a:rPr lang="en-US" sz="' + str(font_size) + '"' + (' b="1"' if bold else '') + '/>'
                '<a:t>' + _sandbox_office_xml_escape(line) + '</a:t></a:r><a:endParaRPr lang="en-US" sz="' + str(font_size) + '"/></a:p>'
            )
        return (
            '<p:sp><p:nvSpPr><p:cNvPr id="' + str(shape_id) + '" name="' + _sandbox_office_xml_escape(name) + '"/>'
            '<p:cNvSpPr txBox="1"/><p:nvPr/></p:nvSpPr><p:spPr><a:xfrm><a:off x="' + str(x) + '" y="' + str(y) + '"/>'
            '<a:ext cx="' + str(cx) + '" cy="' + str(cy) + '"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
            '<a:noFill/><a:ln><a:noFill/></a:ln></p:spPr><p:txBody><a:bodyPr wrap="square"/><a:lstStyle/>'
            + ''.join(paras) + '</p:txBody></p:sp>'
        )

    def slide_xml(slide: dict, idx: int) -> str:
        title = _sandbox_office_text(slide.get('title') or f'Slide {idx}', 500)
        subtitle = _sandbox_office_text(slide.get('subtitle'), 500)
        bullets = slide.get('bullets') if isinstance(slide.get('bullets'), list) else slide.get('items')
        bullet_lines = [_sandbox_office_text(x, 1200) for x in (bullets if isinstance(bullets, list) else [])[:18]]
        if subtitle:
            bullet_lines.insert(0, subtitle)
        shapes = text_shape(2, 'Title', 685800, 450000, 7772400, 760000, [title], 3600, True)
        shapes += text_shape(3, 'Content', 900000, 1350000, 7300000, 4300000, bullet_lines or [''], 2200, False)
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">'
            '<p:cSld><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr>'
            '<p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr>'
            + shapes + '</p:spTree></p:cSld><p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr></p:sld>'
        )

    slide_overrides = ''.join(f'<Override PartName="/ppt/slides/slide{i}.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>' for i in range(1, len(slides) + 1))
    slide_ids = ''.join(f'<p:sldId id="{255 + i}" r:id="rId{i + 1}"/>' for i in range(1, len(slides) + 1))
    pres_rels = [('rId1', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster', 'slideMasters/slideMaster1.xml')]
    pres_rels.extend((f'rId{i + 1}', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide', f'slides/slide{i}.xml') for i in range(1, len(slides) + 1))
    out = io.BytesIO()
    import zipfile
    with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as z:
        z.writestr('[Content_Types].xml', '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/><Override PartName="/ppt/slideMasters/slideMaster1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideMaster+xml"/><Override PartName="/ppt/slideLayouts/slideLayout1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideLayout+xml"/><Override PartName="/ppt/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/><Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/><Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>' + slide_overrides + '</Types>')
        z.writestr('_rels/.rels', rels([('rId1', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument', 'ppt/presentation.xml'), ('rId2', 'http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties', 'docProps/core.xml'), ('rId3', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties', 'docProps/app.xml')]))
        z.writestr('ppt/presentation.xml', '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:sldMasterIdLst><p:sldMasterId id="2147483648" r:id="rId1"/></p:sldMasterIdLst><p:sldIdLst>' + slide_ids + '</p:sldIdLst><p:sldSz cx="9144000" cy="5143500" type="screen16x9"/><p:notesSz cx="6858000" cy="9144000"/></p:presentation>')
        z.writestr('ppt/_rels/presentation.xml.rels', rels(pres_rels))
        z.writestr('ppt/slideMasters/slideMaster1.xml', '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:sldMaster xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:cSld><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr></p:spTree></p:cSld><p:clrMap bg1="lt1" tx1="dk1" bg2="lt2" tx2="dk2" accent1="accent1" accent2="accent2" accent3="accent3" accent4="accent4" accent5="accent5" accent6="accent6" hlink="hlink" folHlink="folHlink"/><p:sldLayoutIdLst><p:sldLayoutId id="2147483649" r:id="rId1"/></p:sldLayoutIdLst><p:txStyles><p:titleStyle/><p:bodyStyle/><p:otherStyle/></p:txStyles></p:sldMaster>')
        z.writestr('ppt/slideMasters/_rels/slideMaster1.xml.rels', rels([('rId1', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout', '../slideLayouts/slideLayout1.xml'), ('rId2', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme', '../theme/theme1.xml')]))
        z.writestr('ppt/slideLayouts/slideLayout1.xml', '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:sldLayout xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" type="blank" preserve="1"><p:cSld name="Blank"><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr></p:spTree></p:cSld><p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr></p:sldLayout>')
        z.writestr('ppt/slideLayouts/_rels/slideLayout1.xml.rels', rels([('rId1', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster', '../slideMasters/slideMaster1.xml')]))
        z.writestr('ppt/theme/theme1.xml', '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="Sandbox"><a:themeElements><a:clrScheme name="Sandbox"><a:dk1><a:srgbClr val="111111"/></a:dk1><a:lt1><a:srgbClr val="FFFFFF"/></a:lt1><a:dk2><a:srgbClr val="333333"/></a:dk2><a:lt2><a:srgbClr val="F2F2F2"/></a:lt2><a:accent1><a:srgbClr val="2563EB"/></a:accent1><a:accent2><a:srgbClr val="10B981"/></a:accent2><a:accent3><a:srgbClr val="F59E0B"/></a:accent3><a:accent4><a:srgbClr val="EF4444"/></a:accent4><a:accent5><a:srgbClr val="8B5CF6"/></a:accent5><a:accent6><a:srgbClr val="06B6D4"/></a:accent6><a:hlink><a:srgbClr val="2563EB"/></a:hlink><a:folHlink><a:srgbClr val="8B5CF6"/></a:folHlink></a:clrScheme><a:fontScheme name="Sandbox"><a:majorFont><a:latin typeface="Aptos Display"/></a:majorFont><a:minorFont><a:latin typeface="Aptos"/></a:minorFont></a:fontScheme><a:fmtScheme name="Sandbox"><a:fillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:fillStyleLst><a:lnStyleLst><a:ln w="9525"><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:ln></a:lnStyleLst><a:effectStyleLst><a:effectStyle/></a:effectStyleLst><a:bgFillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:bgFillStyleLst></a:fmtScheme></a:themeElements></a:theme>')
        z.writestr('docProps/core.xml', '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/"><dc:title>' + _sandbox_office_xml_escape(args.get('title') or 'Presentation') + '</dc:title></cp:coreProperties>')
        z.writestr('docProps/app.xml', '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"><Application>app3 sandbox</Application><Slides>' + str(len(slides)) + '</Slides></Properties>')
        for idx, slide in enumerate(slides, 1):
            z.writestr(f'ppt/slides/slide{idx}.xml', slide_xml(slide, idx))
            z.writestr(f'ppt/slides/_rels/slide{idx}.xml.rels', rels([('rId1', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout', '../slideLayouts/slideLayout1.xml')]))
    raw = out.getvalue()
    with zipfile.ZipFile(io.BytesIO(raw), 'r') as z:
        names = set(z.namelist())
        required = {'[Content_Types].xml', 'ppt/presentation.xml', 'ppt/slides/slide1.xml'}
        if not required.issubset(names):
            raise ValueError('pptx_validation_failed')
    return raw, {'validator': 'minimal_ooxml_zip', 'valid': True, 'slides': len(slides)}


def _sandbox_make_text_office_bytes(fmt: str, args: dict | None = None) -> tuple[bytes, dict]:
    args = dict(args or {})
    import csv
    import html as html_lib
    sections = _sandbox_office_sections(args)
    if fmt == 'csv':
        rows = []
        sheets = args.get('sheets')
        if isinstance(sheets, list) and sheets and isinstance(sheets[0], dict):
            rows = _sandbox_office_rows(sheets[0].get('headers'), sheets[0].get('rows'), max_rows=10000, max_cols=100)
        else:
            table = next((s for s in sections if str(s.get('type') or '').lower() == 'table'), None)
            rows = _sandbox_office_rows((table or {}).get('headers'), (table or {}).get('rows'), max_rows=10000, max_cols=100) if isinstance(table, dict) else [[_sandbox_office_text(args.get('content') or args.get('title') or '', 200000)]]
        sio = io.StringIO()
        writer = csv.writer(sio, lineterminator='\n')
        writer.writerows(rows)
        return sio.getvalue().encode('utf-8-sig'), {'validator': 'csv_writer', 'valid': True, 'rows': len(rows)}
    md_lines = []
    html_parts = []
    for section in sections:
        typ = str(section.get('type') or 'paragraph').lower()
        text = _sandbox_office_text(section.get('text') or section.get('title') or section.get('content'), 50000)
        if typ in {'heading', 'title'}:
            level = max(1, min(int(section.get('level') or 1), 6))
            md_lines.append('#' * level + ' ' + text)
            html_parts.append(f'<h{level}>' + html_lib.escape(text) + f'</h{level}>')
        elif typ in {'list', 'bullets', 'bullet_list'}:
            items = section.get('items') if isinstance(section.get('items'), list) else section.get('bullets')
            html_parts.append('<ul>')
            for item in (items if isinstance(items, list) else []):
                t = _sandbox_office_text(item, 10000)
                md_lines.append('- ' + t)
                html_parts.append('<li>' + html_lib.escape(t) + '</li>')
            html_parts.append('</ul>')
        elif typ == 'table':
            rows = _sandbox_office_rows(section.get('headers'), section.get('rows'), max_rows=500, max_cols=30)
            if rows:
                md_lines.extend([' | '.join(rows[0]), ' | '.join(['---'] * len(rows[0]))])
                md_lines.extend(' | '.join(r) for r in rows[1:])
                html_parts.append('<table>')
                for r in rows:
                    html_parts.append('<tr>' + ''.join('<td>' + html_lib.escape(c) + '</td>' for c in r) + '</tr>')
                html_parts.append('</table>')
        else:
            md_lines.append(text)
            html_parts.append('<p>' + html_lib.escape(text).replace('\n', '<br>') + '</p>')
    if fmt == 'md':
        return ('\n\n'.join(md_lines) + '\n').encode('utf-8'), {'validator': 'utf8_markdown', 'valid': True}
    if fmt == 'html':
        title = html_lib.escape(_sandbox_office_text(args.get('title') or 'Document', 500))
        body = '<!doctype html><html><head><meta charset="utf-8"><title>' + title + '</title></head><body>' + '\n'.join(html_parts) + '</body></html>'
        return body.encode('utf-8'), {'validator': 'utf8_html', 'valid': True}
    if fmt == 'rtf':
        def esc(s: str) -> str:
            return s.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}').replace('\n', '\\par ')
        body = ''.join('\\par ' + esc(line) for line in md_lines)
        return ('{\\rtf1\\ansi\\deff0' + body + '}').encode('utf-8'), {'validator': 'rtf_text', 'valid': True}
    return ('\n\n'.join(md_lines) + '\n').encode('utf-8'), {'validator': 'utf8_text', 'valid': True}


def _sandbox_create_office_file_tool(args: dict | None = None, messages: list | None = None) -> dict:
    if not _sandbox_tools_enabled():
        return {'ok': False, 'error': 'sandbox_tools_disabled'}
    args = dict(args or {})
    ok_limit, limit_error, max_chars = _sandbox_office_limit_args(args)
    if not ok_limit:
        return {'ok': False, 'error': limit_error, 'max_json_chars': max_chars}
    allowed = {'docx', 'xlsx', 'pptx', 'pdf', 'html', 'rtf', 'csv', 'md'}
    raw_path = str(args.get('path') or args.get('filename') or '').strip()
    raw_fmt = str(args.get('format') or '').strip().lower().lstrip('.')
    ext_fmt = os.path.splitext(raw_path)[1].lower().lstrip('.') if raw_path else ''
    fmt = raw_fmt or ext_fmt
    if fmt == 'markdown':
        fmt = 'md'
    if not raw_path:
        return {'ok': False, 'error': 'missing_file_path'}
    if not fmt:
        return {'ok': False, 'error': 'missing_format'}
    if fmt not in allowed:
        return {'ok': False, 'error': 'unsupported_office_format', 'format': fmt, 'supported_formats': sorted(allowed)}
    if ext_fmt and ext_fmt != fmt:
        return {'ok': False, 'error': 'path_extension_format_mismatch', 'path_extension': ext_fmt, 'format': fmt}
    if not ext_fmt:
        raw_path = raw_path.rstrip('.') + '.' + fmt
    try:
        target, rel = _sandbox_resolve_path(raw_path, messages or [])
    except Exception as e:
        return {'ok': False, 'error': str(e or 'invalid_path')}
    makers = {
        'docx': _sandbox_make_docx_bytes,
        'xlsx': _sandbox_make_xlsx_bytes,
        'pptx': _sandbox_make_pptx_bytes,
        'pdf': _sandbox_make_pdf_bytes,
    }
    try:
        if fmt in makers:
            raw, validation = makers[fmt](args)
        else:
            raw, validation = _sandbox_make_text_office_bytes(fmt, args)
        max_bytes = max(1024, min(int(args.get('max_bytes') or app_getenv('SANDBOX_OFFICE_MAX_BYTES', str(80 * 1024 * 1024)) or (80 * 1024 * 1024)), 512 * 1024 * 1024))
        if len(raw) > max_bytes:
            return {'ok': False, 'error': 'office_file_too_large', 'max_bytes': max_bytes, 'size': len(raw)}
        quota_ok, quota_meta = _sandbox_quota_ok(messages or [], incoming_bytes=len(raw), current_path=target, append=False)
        if not quota_ok:
            return quota_meta
        storage_ok, storage_meta = _sandbox_storage_quota_ok(messages or [], incoming_bytes=len(raw), current_path=target, append=False)
        if not storage_ok:
            return storage_meta
        os.makedirs(os.path.dirname(target), exist_ok=True)
        before_snapshot = _sandbox_file_binary_snapshot(target)
        with open(target, 'wb') as f:
            f.write(raw)
        size = int(os.path.getsize(target))
        after_snapshot = _sandbox_file_binary_snapshot(target)
        lineage_meta = {}
        lineage_fn = globals().get('file_context_select_parent_for_output')
        if callable(lineage_fn):
            try:
                lineage_meta = dict(lineage_fn(rel, messages=messages or [], query=str(args.get('title') or args.get('content') or rel)))
            except Exception:
                lineage_meta = {}
        audit = _sandbox_build_binary_audit(rel, before_snapshot, after_snapshot, operation='sandbox_create_office_file', fmt=fmt, lineage=lineage_meta)
        payload = {
            **_sandbox_result_base(messages or []),
            'ok': True,
            '_kind': 'sandbox_office_file',
            'path': rel,
            'format': fmt,
            'size': size,
            'validation': validation,
            'generated_by_assistant': True,
            'artifact_plan': _sandbox_artifact_task_plan_for_messages(args, messages or []),
            'file_context': lineage_meta,
            'parent_file_ids': [str(x or '').strip() for x in (lineage_meta.get('parent_file_ids') or []) if str(x or '').strip()] if isinstance(lineage_meta, dict) else [],
            'source_file_ids': [str(x or '').strip() for x in (lineage_meta.get('source_file_ids') or []) if str(x or '').strip()] if isinstance(lineage_meta, dict) else [],
            'file_edit_audit': audit,
            'edit_audit': audit,
            'publish_instruction': 'Call sandbox_publish_files with this path to deliver a real downloadable file. Do not print /api3/generated-download until sandbox_publish_files returns download_url.',
        }
        auto_publish = globals().get('artifact_manager_auto_publish_generated')
        if callable(auto_publish):
            try:
                payload = dict(auto_publish(payload, args=args, messages=messages or [], source_tool='sandbox_create_office_file') or payload)
            except Exception as publish_exc:
                payload['auto_published'] = False
                payload['publish_error'] = f'{type(publish_exc).__name__}: {publish_exc}'
        return _attach_evidence_ledger_event('sandbox_create_office_file', payload, args)
    except Exception as e:
        return {'ok': False, 'path': rel, 'format': fmt, 'error': f'{type(e).__name__}: {e}'}
